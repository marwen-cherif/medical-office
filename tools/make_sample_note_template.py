"""Genere un modele Word de DEMO « note d'honoraires multi-lignes » (tache 5.1 de
la capability `facturation-multi-lignes`).

Pose `templates/note_multi_lignes_demo.docx` : un tableau avec UNE ligne-modele
porteuse des balises de ligne `<L_*>` (dupliquee a la generation), une ligne de
totaux, un pied de page recapitulatif, les balises de dents agregees
`<DENTS>`/`<NB_DENTS>` et un schema dentaire `<ODONTOGRAMME>` (paragraphe dedie,
capability schema-dentaire-notes). Utilise uniquement python-docx (aucun Word
requis pour CREER le modele ; Word reste requis pour GENERER la note).

Usage (depuis la racine du projet) :
    python -m tools.make_sample_note_template

Pour le tester via le bouton dedie « Note d'honoraires », rangez ensuite ce modele
dans la categorie des notes (Parametrage > Modeles > icone Categorie). Sinon il
apparait dans la generation generique de document.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from crm.templates import templates_dir  # noqa: E402

DEMO_NAME = "note_multi_lignes_demo.docx"


def _runs(paragraph, *parts: str) -> None:
    """Ajoute un run par partie (mime le decoupage de runs de Word : evite que
    prefixe + debut de balise tombent dans un meme run)."""
    for txt in parts:
        paragraph.add_run(txt)


def build(dest: Path) -> Path:
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Note d'honoraires")
    r.bold = True
    r.font.size = Pt(16)

    p = doc.add_paragraph()
    _runs(p, "Patient : ", "<NOM>", " ", "<PRENOM>", "      Emise le : ", "<DATE>")

    # Dents concernees (agregees) : balises document <DENTS> / <NB_DENTS>.
    pd = doc.add_paragraph()
    _runs(pd, "Dents concernees : ", "<DENTS>", "   (", "<NB_DENTS>", " dent(s))")

    doc.add_paragraph()

    # Tableau : entete + 1 ligne-modele <L_*> + ligne de totaux.
    table = doc.add_table(rows=3, cols=6)
    table.style = "Table Grid"  # bordures visibles -> verifier leur conservation
    headers = ["Date", "Acte", "Dents", "Montant", "Regle", "Reste"]
    for cell, txt in zip(table.rows[0].cells, headers):
        cr = cell.paragraphs[0].add_run(txt)
        cr.bold = True

    tags = ["<L_DATE>", "<L_ACTE>", "<L_DENTS>", "<L_MONTANT>", "<L_REGLE>", "<L_RESTE>"]
    for cell, tag in zip(table.rows[1].cells, tags):
        cell.paragraphs[0].add_run(tag)  # balise seule dans la cellule

    totals = table.rows[2].cells
    tr = totals[0].paragraphs[0].add_run("TOTAL")
    tr.bold = True
    totals[3].paragraphs[0].add_run("<TOTAL_DU>")
    totals[4].paragraphs[0].add_run("<TOTAL_REGLE>")
    totals[5].paragraphs[0].add_run("<RESTE_A_PAYER>")

    doc.add_paragraph()
    foot = doc.add_paragraph()
    _runs(foot, "Nombre d'actes : ", "<NB_ACTES>", "      Reste a payer : ", "<RESTE_A_PAYER>")

    # Schema dentaire (image) : balise <ODONTOGRAMME> SEULE dans son paragraphe dedie
    # (consigne : evite tout debordement ; remplacee a la generation par le schema
    # anatomique des dents concernees). Capability schema-dentaire-notes.
    doc.add_paragraph()
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run("Schema des dents concernees").italic = True
    schema = doc.add_paragraph()
    schema.alignment = WD_ALIGN_PARAGRAPH.CENTER
    schema.add_run("<ODONTOGRAMME>")  # balise seule -> image inseree a sa place

    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dest))
    return dest


def main() -> None:
    dest = build(templates_dir() / DEMO_NAME)
    print(f"Modele de demo cree : {dest}")


if __name__ == "__main__":
    main()
