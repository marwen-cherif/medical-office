"""Gestion des modeles de documents (.docx) du cabinet.

Un dossier `templates/` contient un .docx par type de document (note
d'honoraires, devis, recu...). L'app peut les lister, en creer un nouveau et
les ouvrir dans Word pour edition.
"""

from __future__ import annotations

import os
import shutil
import subprocess
import sys
from dataclasses import dataclass
from pathlib import Path

from .db import app_dir


def templates_dir() -> Path:
    d = app_dir() / "templates"
    d.mkdir(parents=True, exist_ok=True)
    return d


@dataclass
class Template:
    name: str          # nom logique (sans extension), ex. "note_honoraires"
    path: Path

    @property
    def label(self) -> str:
        return self.name.replace("_", " ").strip().capitalize()


def list_templates() -> list[Template]:
    out: list[Template] = []
    for p in sorted(templates_dir().glob("*.docx")):
        if p.name.startswith("~$"):  # fichiers temporaires Word ouverts
            continue
        out.append(Template(name=p.stem, path=p))
    return out


def get_template(name: str) -> Template | None:
    p = templates_dir() / f"{name}.docx"
    return Template(name=name, path=p) if p.exists() else None


def create_template(name: str, copy_from: Path | None = None) -> Template:
    """Cree un modele. Copie un .docx existant si fourni, sinon un docx minimal."""
    safe = _safe_name(name)
    dest = templates_dir() / f"{safe}.docx"
    if dest.exists():
        raise FileExistsError(f"Le modele '{safe}' existe deja.")
    if copy_from and copy_from.exists():
        shutil.copyfile(copy_from, dest)
    else:
        _write_blank_docx(dest)
    return Template(name=safe, path=dest)


def rename_template(template: Template, new_name: str) -> Template:
    """Renomme un modele. Leve FileExistsError si la cible existe deja."""
    safe = _safe_name(new_name)
    if safe == template.name:
        return template
    dest = templates_dir() / f"{safe}.docx"
    if dest.exists():
        raise FileExistsError(f"Le modele '{safe}' existe deja.")
    template.path.rename(dest)
    return Template(name=safe, path=dest)


def delete_template(template: Template) -> None:
    """Supprime le fichier .docx du modele."""
    if template.path.exists():
        template.path.unlink()


def open_in_word(template: Template) -> None:
    """Ouvre le modele dans l'application par defaut (Word) pour edition."""
    path = str(template.path)
    if os.name == "nt":
        os.startfile(path)  # type: ignore[attr-defined]
    elif sys.platform == "darwin":
        subprocess.Popen(["open", path])
    else:
        subprocess.Popen(["xdg-open", path])


def _safe_name(name: str) -> str:
    cleaned = "".join(c if c.isalnum() else "_" for c in name.strip().lower())
    while "__" in cleaned:
        cleaned = cleaned.replace("__", "_")
    return cleaned.strip("_") or "modele"


def _write_blank_docx(dest: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_paragraph("Cabinet Dr Aslem Gouiaa")
    doc.add_paragraph("")
    doc.add_paragraph("Patient : <NOM> <PRENOM>")
    doc.add_paragraph("Date : <DATE>")
    doc.add_paragraph("Acte : <ACTE>")
    doc.add_paragraph("Montant : <MONTANT>")
    doc.add_paragraph("")
    doc.add_paragraph(
        "Astuce : utilisez les balises <NOM>, <PRENOM>, <DATE>, <ACTE>, <MONTANT>."
    )
    doc.save(str(dest))
