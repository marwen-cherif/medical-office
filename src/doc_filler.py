from __future__ import annotations

import re
import shutil
import tempfile
from contextlib import contextmanager
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Emu

# wdFormatPDF constant
WD_FORMAT_PDF = 17

# Balises *image* : remplacees par une image en ligne (et non par du texte). Une
# balise image presente dans un modele mais non fournie est simplement videe (pas de
# balise litterale dans le rendu). Cf. capability schema-dentaire-notes (<ODONTOGRAMME>).
IMAGE_TAGS = {"ODONTOGRAMME"}

# Boite max d'une image inseree (l'image est mise a l'echelle en preservant son ratio
# pour tenir dedans, quel que soit son format portrait/paysage). Un schema dentaire est
# portrait : la hauteur est en general le facteur limitant.
_IMG_MAX_W_CM = 9.5
_IMG_MAX_H_CM = 13.0


class WordSession:
    """Context manager autour d'une instance Word COM unique."""

    def __init__(self):
        self._word = None
        self._pythoncom = None

    def __enter__(self):
        import pythoncom
        import win32com.client

        self._pythoncom = pythoncom
        pythoncom.CoInitialize()
        self._word = win32com.client.DispatchEx("Word.Application")
        self._word.Visible = False
        self._word.DisplayAlerts = 0
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        try:
            if self._word is not None:
                self._word.Quit(SaveChanges=0)
        finally:
            self._word = None
            if self._pythoncom is not None:
                self._pythoncom.CoUninitialize()
                self._pythoncom = None

    def fill_and_export_pdf(
        self,
        template_path: Path,
        replacements: dict[str, str],
        pdf_output: Path,
        line_rows: list[dict[str, str]] | None = None,
        images: dict[str, Path] | None = None,
    ) -> Path:
        """Remplace les placeholders via python-docx, exporte en PDF via Word COM.

        `line_rows` (optionnel) : note multi-lignes -> la ligne-modele `<L_*>` est
        repetee une fois par entree (cf. `_fill_docx`/`expand_table_rows`).
        `images` (optionnel) : map balise->chemin image (sans chevrons, ex.
        `{"ODONTOGRAMME": Path(...)}`) ; chaque balise image est remplacee par
        l'image en ligne (cf. `_insert_images`).
        """
        pdf_output.parent.mkdir(parents=True, exist_ok=True)

        with tempfile.TemporaryDirectory() as tmp:
            filled_docx = Path(tmp) / "filled.docx"
            _fill_docx(template_path, replacements, filled_docx,
                       line_rows=line_rows, images=images)

            doc = self._word.Documents.Open(str(filled_docx), ReadOnly=False)
            try:
                doc.SaveAs(str(pdf_output), FileFormat=WD_FORMAT_PDF)
            finally:
                doc.Close(SaveChanges=0)

        return pdf_output


_PLACEHOLDER_RE = re.compile(r"<([A-Z0-9_]+)>")


def _p_elem_text(p_elem) -> str:
    """Texte concatene d'un paragraphe (element w:p), tous runs confondus."""
    return "".join(t.text or "" for t in p_elem.iter(qn("w:t")))


def extract_placeholders(template_path: Path) -> list[str]:
    """Liste les balises <TAG> presentes dans le .docx, dans l'ordre, sans doublon.

    Parcourt le corps, les zones de texte, et les en-tetes/pieds (meme traversee que
    `_fill_docx`). Une balise peut etre eclatee sur plusieurs runs : on reconstruit donc
    le texte complet de chaque paragraphe avant la recherche.
    """
    doc = Document(str(template_path))
    chunks: list[str] = []

    for para in doc.paragraphs:
        chunks.append(para.text)
    for txbx in doc.element.body.iter(qn("w:txbxContent")):
        for p_elem in txbx.iter(qn("w:p")):
            chunks.append(_p_elem_text(p_elem))
    # Cellules de tableau : `doc.paragraphs` ne les traverse pas. Indispensable pour
    # detecter les balises de ligne `<L_*>` (qui vivent dans la ligne-modele) et les
    # balises document placees dans un tableau (ex. ligne de totaux <TOTAL_DU>).
    for tbl in doc.element.body.iter(qn("w:tbl")):
        for p_elem in tbl.iter(qn("w:p")):
            chunks.append(_p_elem_text(p_elem))
    for section in doc.sections:
        for para in section.header.paragraphs:
            chunks.append(para.text)
        for para in section.footer.paragraphs:
            chunks.append(para.text)

    seen: dict[str, None] = {}
    for text in chunks:
        for match in _PLACEHOLDER_RE.findall(text):
            seen.setdefault(match, None)
    return list(seen.keys())


# --- Note multi-lignes : convention de balises ligne/document -----------------
# Une balise de ligne porte le prefixe `L_` ; elle est repetee une fois par ligne
# de la note (cf. capability facturation-multi-lignes). Les autres balises sont
# des balises *document*, remplies une seule fois. La distinction est purement
# additive : la regex `<([A-Z0-9_]+)>` matche deja `L_*`.
LINE_TAG_PREFIX = "L_"


def is_line_tag(tag: str) -> bool:
    """Vrai si `tag` (sans chevrons) est une balise de ligne (prefixe `L_`)."""
    return tag.upper().startswith(LINE_TAG_PREFIX)


def line_tag_column(tag: str) -> str:
    """Nom de colonne d'une balise de ligne : balise sans prefixe, en minuscules
    (`L_MONTANT` -> `montant`). Sans effet sur une balise document."""
    if is_line_tag(tag):
        return tag[len(LINE_TAG_PREFIX):].lower()
    return tag.lower()


def classify_placeholders(template_path: Path) -> tuple[list[str], list[str]]:
    """Renvoie `(balises_document, balises_de_ligne)` du modele, dans l'ordre.

    `extract_placeholders` reste la liste a plat (compat consommateurs) ; cette
    fonction la classe par convention de prefixe. Un modele est une « note
    multi-lignes » des qu'il porte >= 1 balise de ligne `L_*`.
    """
    tags = extract_placeholders(template_path)
    doc_tags = [t for t in tags if not is_line_tag(t)]
    line_tags = [t for t in tags if is_line_tag(t)]
    return doc_tags, line_tags


def _find_template_rows(scope) -> list:
    """Lignes de tableau (`w:tr`) portant >= 1 balise de ligne `L_*`.

    Reconstruit le texte complet de chaque ligne (balises eventuellement eclatees
    sur plusieurs runs) avant la recherche. `scope` est un element lxml (corps).
    """
    rows = []
    for tr in scope.iter(qn("w:tr")):
        text = "".join(t.text or "" for t in tr.iter(qn("w:t")))
        if any(is_line_tag(m) for m in _PLACEHOLDER_RE.findall(text)):
            rows.append(tr)
    return rows


def _validate_template_row(tr) -> None:
    """Refuse une ligne-modele a la structure non supportee (D3) : message clair."""
    if tr.find(qn("w:tbl")) is not None or next(tr.iter(qn("w:tbl")), None) is not None:
        raise ValueError(
            "Modele « note multi-lignes » non supporte : tableau imbrique dans la "
            "ligne-modele <L_*>."
        )
    if next(tr.iter(qn("w:vMerge")), None) is not None:
        raise ValueError(
            "Modele « note multi-lignes » non supporte : cellules fusionnees "
            "verticalement dans la ligne-modele <L_*>."
        )


def expand_table_rows(doc, line_rows: list[dict[str, str]], template_rows=None) -> None:
    """Repete la ligne-modele d'un tableau Word : une copie par entree de
    `line_rows`, remplie via le `_replace_in_para_elem` *reutilise*, puis retire
    la ligne-modele d'origine (retrait simple si 0 ligne).

    Fonction **additive** (D3) : sans ligne-modele `<L_*>`, ne fait rien. La
    duplication est purement structurelle (`deepcopy` du `w:tr` complet -> mise en
    forme `w:trPr`/`w:tcPr`/`w:rPr` preservee) ; le remplissage des cellules
    reutilise la logique de redistribution de runs sans la modifier.

    Chaque entree de `line_rows` est un dict `{ "<L_DATE>": "...", ... }` (balises
    de ligne completes, chevrons inclus).
    """
    rows = template_rows if template_rows is not None else _find_template_rows(doc.element.body)
    if not rows:
        return
    if len(rows) > 1:
        raise ValueError(
            "Modele « note multi-lignes » non supporte : plusieurs lignes-modeles "
            "<L_*> detectees. Une seule ligne de tableau doit porter les balises <L_*>."
        )
    tr = rows[0]
    _validate_template_row(tr)
    parent = tr.getparent()
    idx = parent.index(tr)
    for offset, row_repl in enumerate(line_rows):
        clone = deepcopy(tr)
        for p_elem in clone.iter(qn("w:p")):
            _replace_in_para_elem(p_elem, row_repl)
        parent.insert(idx + 1 + offset, clone)
    parent.remove(tr)


# --- Insertion d'images en remplacement d'une balise (schema-dentaire-notes) -----
# Capacite additive : une balise document (ex. <ODONTOGRAMME>) est remplacee par une
# image en ligne, a son emplacement, en gerant les balises eclatees sur plusieurs runs
# Word. Couvre corps, cellules de tableau et en-tetes/pieds (parcours python-docx, ou
# l'on dispose d'objets Run pour `add_picture`). Une balise seule dans son paragraphe
# (usage recommande) est le cas nominal.


def _image_size_emu(path: Path, max_w_emu: int, max_h_emu: int) -> tuple[int, int]:
    """Dimensions (EMU) d'une image mise a l'echelle pour tenir dans la boite max en
    preservant son ratio. Pillow (deja embarque) lit les dimensions en pixels."""
    from PIL import Image

    with Image.open(str(path)) as im:
        pw, ph = im.size
    if pw <= 0 or ph <= 0:
        return max_w_emu, max_h_emu
    scale = min(max_w_emu / pw, max_h_emu / ph)
    return Emu(int(pw * scale)), Emu(int(ph * scale))


def _walk_paragraphs(container):
    """Tous les paragraphes d'un conteneur python-docx (corps, cellule, en-tete/pied),
    en descendant recursivement dans les cellules de tableau."""
    yield from container.paragraphs
    for tbl in container.tables:
        for row in tbl.rows:
            for cell in row.cells:
                yield from _walk_paragraphs(cell)


def _insert_image_in_paragraph(paragraph, sized: dict[str, tuple[Path, int, int]]) -> None:
    """Remplace, dans `paragraph`, la 1re occurrence de chaque balise image par son
    image en ligne. Localise la balise meme eclatee sur plusieurs runs (meme principe
    que `_replace_in_para_elem`), preserve le texte avant/apres, et insere l'image
    dans le run de debut (l'ordre prefixe / image / suffixe est conserve)."""
    runs = paragraph.runs
    if not runs:
        return
    texts = [r.text or "" for r in runs]
    full = "".join(texts)
    for tag, (path, w_emu, h_emu) in sized.items():
        ph = f"<{tag}>"
        start = full.find(ph)
        if start < 0:
            continue
        end = start + len(ph)
        char_run: list[int] = []
        for i, t in enumerate(texts):
            char_run.extend([i] * len(t))
        first = char_run[start]
        last = char_run[end - 1]
        first_run_start = sum(len(texts[i]) for i in range(first))
        prefix = texts[first][: start - first_run_start]
        last_run_start = sum(len(texts[i]) for i in range(last))
        suffix = texts[last][end - last_run_start:]

        runs[first].text = prefix
        for i in range(first + 1, last + 1):
            runs[i].text = suffix if i == last else ""
        runs[first].add_picture(str(path), width=w_emu, height=h_emu)
        if first == last and suffix:
            runs[first].add_text(suffix)
        return  # une balise image (document) au plus par paragraphe


def _insert_images(doc, images: dict[str, Path]) -> None:
    """Insere chaque image fournie a l'emplacement de sa balise, dans tout le document
    (corps, tableaux, en-tetes/pieds, et **zones de texte** du corps — les modeles « lettre »
    placent souvent le corps dans une zone de texte). Les balises image non fournies ont
    deja ete videes par le remplissage texte (cf. `_fill_docx`)."""
    from docx.text.paragraph import Paragraph

    max_w, max_h = Cm(_IMG_MAX_W_CM), Cm(_IMG_MAX_H_CM)
    sized = {
        tag: (path, *_image_size_emu(path, max_w, max_h)) for tag, path in images.items()
    }
    containers = [doc]
    for section in doc.sections:
        containers.append(section.header)
        containers.append(section.footer)
    for container in containers:
        for para in _walk_paragraphs(container):
            _insert_image_in_paragraph(para, sized)

    # Zones de texte du corps (w:txbxContent) : non traversees par python-docx
    # (`doc.paragraphs`). On enveloppe chaque w:p en Paragraph rattache au document
    # (sa `part` porte l'image inseree via add_picture). Couvre aussi les zones de texte
    # nichees dans un tableau du corps (iter recursif).
    for txbx in doc.element.body.iter(qn("w:txbxContent")):
        for p_elem in txbx.iter(qn("w:p")):
            _insert_image_in_paragraph(Paragraph(p_elem, doc), sized)


def _fill_docx(
    template_path: Path,
    replacements: dict[str, str],
    output_path: Path,
    line_rows: list[dict[str, str]] | None = None,
    images: dict[str, Path] | None = None,
) -> None:
    """Remplace les placeholders du docx (corps, zones de texte, en-tetes/pieds,
    cellules de tableau) et, si `line_rows` est fourni, repete la ligne-modele.
    Si `images` est fourni, remplace les balises image par l'image correspondante ;
    toute balise image NON fournie est videe (jamais rendue litteralement).

    `line_rows=None` => comportement historique mono-valeur (aucune expansion).
    `images=None` => aucune insertion d'image (comportement historique).
    """
    images = images or {}
    # Vide les balises image non fournies (ensemble vide de dents, etc.) pour qu'aucune
    # balise litterale ne subsiste ; les balises image fournies restent litterales (la
    # passe d'insertion les localise apres le remplissage texte). Copie defensive.
    replacements = dict(replacements)
    for tag in IMAGE_TAGS:
        if tag not in images:
            replacements.setdefault(f"<{tag}>", "")

    doc = Document(str(template_path))

    # Ligne(s)-modele : reperees une fois pour les exclure de la passe document
    # (elles sont traitees par l'expansion) et alimenter `expand_table_rows`.
    template_rows = _find_template_rows(doc.element.body)
    template_row_ids = {id(tr) for tr in template_rows}

    # Corps principal
    for para in doc.paragraphs:
        _replace_in_para(para, replacements)

    # Zones de texte (w:txbxContent)
    for txbx in doc.element.body.iter(qn("w:txbxContent")):
        for p_elem in txbx.iter(qn("w:p")):
            _replace_in_para_elem(p_elem, replacements)

    # En-têtes et pieds de page
    for section in doc.sections:
        for para in section.header.paragraphs:
            _replace_in_para(para, replacements)
        for para in section.footer.paragraphs:
            _replace_in_para(para, replacements)

    # Cellules de tableau (balises document) : `doc.paragraphs` ne les traverse
    # pas. On remplit ici les balises document presentes dans les tableaux (ex.
    # ligne de totaux <TOTAL_DU>), en excluant la/les ligne(s)-modele.
    for tbl in doc.element.body.iter(qn("w:tbl")):
        for tr in tbl.iter(qn("w:tr")):
            if id(tr) in template_row_ids:
                continue
            for p_elem in tr.iter(qn("w:p")):
                _replace_in_para_elem(p_elem, replacements)

    # Expansion de la ligne-modele (note multi-lignes) si demandee.
    if line_rows is not None:
        expand_table_rows(doc, line_rows, template_rows=template_rows)

    # Insertion des images (apres le texte et l'expansion) : remplace chaque balise
    # image fournie par l'image en ligne, a son emplacement.
    if images:
        _insert_images(doc, images)

    doc.save(str(output_path))


def _replace_in_para(para, replacements: dict[str, str]) -> None:
    _replace_in_para_elem(para._element, replacements)


_XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"


def _replace_in_para_elem(p_elem, replacements: dict[str, str]) -> None:
    runs = list(p_elem.iter(qn("w:r")))
    if not runs:
        return

    # Collect (run, w:t element, current text) for each run
    run_data = []
    for r in runs:
        t = r.find(qn("w:t"))
        run_data.append([r, t, t.text if (t is not None and t.text) else ""])

    full_text = "".join(rd[2] for rd in run_data)
    new_full = full_text
    for ph, val in replacements.items():
        new_full = new_full.replace(ph, val)

    if new_full == full_text:
        return

    # Redistribute new_full back to runs.
    # Build a character → run-index map for the original text.
    char_run = []
    for i, rd in enumerate(run_data):
        char_run.extend([i] * len(rd[2]))

    # For each placeholder occurrence, find the span of runs it covers,
    # assign the replacement value to the first run of that span and
    # blank out the others. Rebuild new run texts in one pass.
    new_run_texts = [rd[2] for rd in run_data]

    for ph, val in replacements.items():
        ph_len = len(ph)
        start = 0
        while True:
            idx = "".join(new_run_texts).find(ph, start)
            if idx < 0:
                break
            end = idx + ph_len

            # Rebuild a fresh char_run from current new_run_texts
            cr = []
            for i, t in enumerate(new_run_texts):
                cr.extend([i] * len(t))

            first = cr[idx] if idx < len(cr) else None
            last = cr[end - 1] if end - 1 < len(cr) else None
            if first is None or last is None:
                break

            # Compute prefix and suffix within the first/last runs
            first_run_start = sum(len(new_run_texts[i]) for i in range(first))
            prefix = new_run_texts[first][: idx - first_run_start]

            last_run_start = sum(len(new_run_texts[i]) for i in range(last))
            suffix = new_run_texts[last][end - last_run_start :]

            # Use the run containing the middle character of the placeholder
            # so the replacement inherits the dominant formatting (e.g. bold)
            mid_char = idx + ph_len // 2
            mid_run = cr[mid_char] if mid_char < len(cr) else first

            # Empty all runs before mid_run (they held the start of the placeholder)
            new_run_texts[first] = prefix
            for i in range(first + 1, mid_run):
                new_run_texts[i] = ""
            # mid_run is always entirely inside the placeholder span, so no prefix
            new_run_texts[mid_run] = val + (suffix if mid_run == last else "")
            # Empty runs after mid_run up to (but not including) last
            for i in range(mid_run + 1, last + 1):
                new_run_texts[i] = suffix if i == last else ""

            start = idx + len(val)

    # Write back to runs (converting newlines into <w:br/> line breaks)
    for rd, new_text in zip(run_data, new_run_texts):
        r, t_elem, old_text = rd
        if new_text == old_text:
            continue
        _set_run_text(r, t_elem, new_text)


def _set_run_text(r, t_elem, text: str) -> None:
    """Ecrit `text` dans le run `r`, en transformant les retours a la ligne
    (\\n) en elements <w:br/> pour qu'ils soient rendus par Word."""
    from lxml import etree

    text = text.replace("\r\n", "\n").replace("\r", "\n")

    if "\n" not in text:
        if t_elem is None:
            t_elem = etree.SubElement(r, qn("w:t"))
        t_elem.text = text
        if text and text != text.strip():
            t_elem.set(_XML_SPACE, "preserve")
        elif _XML_SPACE in t_elem.attrib and text == text.strip():
            del t_elem.attrib[_XML_SPACE]
        return

    # Multi-ligne : on reconstruit le contenu du run avec des <w:br/>.
    # On retire les anciens noeuds de texte/saut, on garde <w:rPr> (le style).
    for child in list(r):
        if child.tag in (qn("w:t"), qn("w:br"), qn("w:cr")):
            r.remove(child)

    for i, line in enumerate(text.split("\n")):
        if i > 0:
            etree.SubElement(r, qn("w:br"))
        t = etree.SubElement(r, qn("w:t"))
        t.text = line
        t.set(_XML_SPACE, "preserve")


def format_montant(value: float) -> str:
    return f"{value:,.3f}".replace(",", " ").replace(".", ",")
